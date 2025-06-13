import os
import requests
from bs4 import BeautifulSoup
from docx import Document
import re
from urllib.parse import urljoin

def sanitize_filename(name):
    """
    Membersihkan string agar menjadi nama file yang valid.
    """
    cleaned_name = re.sub(r'[\\/*?:"<>|]', "", name)
    cleaned_name = re.sub(r'\s+', ' ', cleaned_name).strip()
    return cleaned_name[:100]

def process_urls_from_file(input_filename):
    """
    Membaca daftar URL dari file, menyalin kontennya termasuk link (href),
    dan menyimpannya ke file Word dengan nama sesuai judul halaman.
    """
    if not os.path.exists(input_filename):
        print(f"Error: File '{input_filename}' tidak ditemukan!")
        print("Silakan buat file tersebut dan isi dengan daftar URL (satu per baris).")
        return

    output_folder = "Hasil Salinan Web (dengan Link)"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Folder '{output_folder}' dibuat untuk menyimpan hasil.")

    with open(input_filename, 'r') as file:
        urls = file.readlines()

    print(f"Ditemukan {len(urls)} URL untuk diproses...")

    for i, base_url in enumerate(urls):
        base_url = base_url.strip()
        if not base_url:
            continue

        print(f"\n({i+1}/{len(urls)}) Memproses URL: {base_url}")

        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(base_url, headers=headers, timeout=15)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # --- BAGIAN BARU: Proses semua link (tag <a>) ---
            # Cari semua tag 'a' yang memiliki atribut 'href'
            for link_tag in soup.find_all('a', href=True):
                link_text = link_tag.get_text(strip=True)
                
                # Jangan proses link jika tidak ada teksnya (misal, link pada gambar)
                if not link_text:
                    continue
                    
                # Ambil nilai href
                href = link_tag.get('href')
                
                # Ubah URL relatif (misal: /profil) menjadi URL absolut
                absolute_href = urljoin(base_url, href)
                
                # Buat string baru untuk menggantikan tag <a>
                replacement_string = f"{link_text} ({absolute_href})"
                
                # Gantikan tag <a> dengan string yang sudah diformat
                link_tag.replace_with(replacement_string)
            # --- AKHIR BAGIAN BARU ---

            page_title = soup.title.string if soup.title else f"Tanpa Judul - {i+1}"
            doc_filename = sanitize_filename(page_title) + ".docx"
            full_path = os.path.join(output_folder, doc_filename)

            document = Document()
            document.add_heading(page_title, level=1)
            document.add_paragraph(f"Konten ini disalin dari: {base_url}\n")

            body_content = soup.find('body')
            if body_content:
                # Sekarang get_text() akan menyertakan string link yang sudah kita format
                all_text = body_content.get_text(separator='\n', strip=True)
                document.add_paragraph(all_text)
            else:
                document.add_paragraph("Konten utama (body) tidak dapat ditemukan.")

            document.save(full_path)
            print(f"✅ Berhasil disimpan ke: '{full_path}'")

        except requests.exceptions.RequestException as e:
            print(f"❌ Gagal memproses URL {base_url}. Error: {e}")
        except Exception as e:
            print(f"❌ Terjadi error tak terduga saat memproses {base_url}. Error: {e}")

if __name__ == "__main__":
    nama_file_sumber = "convert_web_document/urls.txt"
    process_urls_from_file(nama_file_sumber)
    print("\n--- Selesai ---")